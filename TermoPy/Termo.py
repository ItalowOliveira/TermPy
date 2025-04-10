import pandas as pd
from docx import Document
from datetime import date
import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog

data_atual = date.today()

def somente_numeros(char):
    return char.isdigit()

def somente_letras(char):
    return char.isalpha()

def carregar_contador():
        try:
            with open('contador.txt', 'r') as f:
                return int(f.read())
        except FileNotFoundError:
            return 0  

def salvar_contador(contador):
        with open('contador.txt', 'w') as f:
            f.write(str(contador))

def incrementar_contador():
        contador = carregar_contador()
        contador += 1
        salvar_contador(contador)
        return contador


def RetornaDados():
    CPF = Edit2.get()
    Nome = Edit1.get()
    CodEquipamento = ''
    NumeroSequencial = incrementar_contador()
    Linha = 1
    Coluna = 0

    data_atual = date.today()
    data_em_texto = "0{}/0{}/{}".format(data_atual.day, data_atual.month, data_atual.year)

    doc = Document("Termo.docx")
    tabela = doc.tables[0]

    for paragrafo in doc.paragraphs:
        if 'NumeroTermo' in paragrafo.text:
         paragrafo.text = paragrafo.text.replace('NumeroTermo', str(NumeroSequencial))

    for paragrafo in doc.paragraphs:
        if 'NomeColaborador' in paragrafo.text:
         paragrafo.text = paragrafo.text.replace('NomeColaborador', Nome)

    for paragrafo in doc.paragraphs:
        if 'CPFColaborador' in paragrafo.text:
         paragrafo.text = paragrafo.text.replace('CPFColaborador', CPF)
    
    caminho_arquivo = "equipamentos.xlsx"
    op1 = variavel_selecao.get()

    if op1 == 'Computador':
        CodEquipamento = 'AFPC'+Edit3.get()
        df = pd.read_excel(caminho_arquivo, sheet_name="Computadores", header=1)
        selected_columns = ['NOME', 'MODELO', 'N/S', 'PREÇO']
        linha = df.loc[df['NOME'] == CodEquipamento, selected_columns]

        if not linha.empty:
            nome = linha.iloc[0]['NOME']
            modelo = linha.iloc[0]['MODELO']
            numero_serie = linha.iloc[0]['N/S']
            preco = linha.iloc[0]['PREÇO']
            infoComputadores = (nome, modelo, numero_serie, preco)
            Lbl5.config(text=infoComputadores)
            tabela.cell(Linha, Coluna).text = str(data_em_texto)
            Coluna += 2
            tabela.cell(Linha, Coluna).text = str(nome)
            Coluna += 1
            tabela.cell(Linha, Coluna).text = "Notebook"
            Coluna += 1
            tabela.cell(Linha, Coluna).text = str(modelo)
            Coluna += 1     
            tabela.cell(Linha, Coluna).text = str(numero_serie)
            Coluna += 1
            tabela.cell(Linha, Coluna).text = str(preco)
            Coluna += 1
            tabela.cell(Linha, Coluna).text = "Teclado + Mouse"

    elif op1 == 'Equipamento':
        CodEquipamento = 'AFEQPTO'+Edit3.get() 
        df = pd.read_excel(caminho_arquivo, sheet_name="Equipamentos", header=1)
        selected_columns = ['NOME', 'MODELO', 'N/S', 'PREÇO']
        linha = df.loc[df['NOME'] == CodEquipamento, selected_columns]

        if not linha.empty:
            nome = linha.iloc[0]['NOME']
            modelo = linha.iloc[0]['MODELO']
            numero_serie = linha.iloc[0]['N/S']
            preco = linha.iloc[0]['PREÇO']
            infoEquipamento = (nome, modelo, numero_serie, preco)
            Lbl5.config(text=infoEquipamento)
            tabela.cell(Linha, Coluna).text = str(nome)
            Coluna += 1
            tabela.cell(Linha, Coluna).text = str(modelo)
            Coluna += 1     
            tabela.cell(Linha, Coluna).text = str(numero_serie)
            Coluna += 1
            tabela.cell(Linha, Coluna).text = str(preco)


    elif op1 == 'Celulares':
        CodEquipamento = 'AFCEL'+Edit3.get()
        df = pd.read_excel(caminho_arquivo, sheet_name="Celulares", header=1)
        selected_columns = ['NOME', 'DESCRIÇÃO', 'MODELO', 'N/S', 'PREÇO']
        linha = df.loc[df['NOME'] == CodEquipamento, selected_columns]
        if not linha.empty:
            nome = linha.iloc[0]['NOME']
            modelo = linha.iloc[0]['MODELO']
            numero_serie = linha.iloc[0]['N/S']
            preco = linha.iloc[0]['PREÇO']
            infoCelulares = (nome, modelo, numero_serie, preco)
            Lbl5.config(text=infoCelulares)
            tabela.cell(Linha, Coluna).text = str(nome)
            Coluna += 1
            tabela.cell(Linha, Coluna).text = str(modelo)
            Coluna += 1     
            tabela.cell(Linha, Coluna).text = str(numero_serie)
            Coluna += 1
            tabela.cell(Linha, Coluna).text = str(preco)
            Coluna += 1
            tabela.cell(Linha, Coluna).text = "Carregador"

    else:
            print('Não escolheu')


def SalvarTermo():
    RetornaDados()
    doc = Document("Termo.docx")
    caminho_arquivo = "equipamentos.xlsx"
    NumeroSequencial = incrementar_contador()

    nomeArquivo = (f"TermoDeResponsabilidadeNº{NumeroSequencial}")
    caminhoSavearquivo = filedialog.asksaveasfilename(
    defaultextension=".docx",
    filetypes=[("Documentos do Word", "*.docx")],
    title="Salvar como",
    initialfile=nomeArquivo
    )

    if caminho_arquivo:
        doc.save(caminhoSavearquivo)
        text1 = (f"Arquivo salvo em: {caminhoSavearquivo}")
        messagebox.showwarning("Arquivo Salvo", text1)
    else:
        text2 = ("Operação de salvamento cancelada.")
        messagebox.showwarning("Erro salvamento", text1)

    doc.save('Termo.Editado.docx')
    return


def ValidacaoDados():

    nomeColaborador = Edit1.get()
    cpfColaborador = Edit2.get()
    numeroEquipamento = Edit3.get()

    if nomeColaborador == '' or cpfColaborador == '' or variavel_selecao.get() == "Escolha uma opção" or numeroEquipamento == '':
         textotest = "Preencha todos os campos pra prosseguir."
         messagebox.showwarning("Erro Preenchimento", textotest)
    else:
        RetornaDados()
    return

root = tk.Tk()
root.title("Gerador de Termo")
#root.resizable(width=False, height=False)
root.geometry("400x300") 

validadorNumeros = root.register(somente_numeros)
validadorLetras = root.register(somente_letras)
Lbl1 = tk.Label(root, text='Nome do Colaborador')
Lbl1.pack()

Edit1 = tk.Entry(root)
Edit1.pack()
Edit1.focus_set()

Lbl2 = tk.Label(root, text='CPF:')
Lbl2.pack()

Edit2 = tk.Entry(root,  validate="key", validatecommand=(validadorNumeros, "%S"))
Edit2.pack()


Lbl4= tk.Label(root)

Lbl3 = tk.Label(root, text='Tipo de Equipamento:')
Lbl3.pack()

variavel_selecao = tk.StringVar()
variavel_selecao.set("Escolha uma opção")

opcoes = ["Computador", "Equipamento", "Celulares"]

menu_select = tk.OptionMenu(root, variavel_selecao, *opcoes)
menu_select.pack()

Lbl4 = tk.Label(root, text='Cod do Item')
Lbl4.pack()


Edit3 = tk.Entry(root, validate="key", validatecommand=(validadorNumeros, "%S"))
Edit3.pack()

SubmitBtn = tk.Button(root,text="Adicionar Termo", command=ValidacaoDados)
SubmitBtn .pack()

SubmitBtn = tk.Button(root,text="Gerar Termo", command=SalvarTermo)
SubmitBtn .pack()

Lbl5 = tk.Label(root, text="Teste")
Lbl5.pack()

Lbl6 = tk.Label(root, text="Teste")
Lbl6.pack()

Lbl7 = tk.Label(root, text="Teste")
Lbl7.pack()

Lbl8 = tk.Label(root, text="Teste")
Lbl8.pack()


root.mainloop()